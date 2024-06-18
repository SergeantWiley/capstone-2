from docx import Document

# Create a new Document
doc = Document()
'''
THIS DOCUMENT WAS IMPROVED AND MODIFIED BY AI. ORIGINAL REPORT PRIOR MODIFICATION IS INCLUDED AS WELL

'''
# Title
doc.add_heading('Capstone Project: Analysis of Riot Direct in VALORANT', 0)

# Introduction Section
doc.add_heading('Introduction', level=1)
intro_text = ("VALORANT is a 5v5 FPS game developed by Riot Games. Riot Games is well known for their innovative technical "
              "expertise in backend technology. One of the most notable examples is Riot Direct. Riot Direct is a physical "
              "direct connection to Riot’s servers. By having a designated internet cable from common Internet Service Providers, "
              "it improves the gaming experience. This seemingly small feature not known by many players has massive implications "
              "for game performance. In this report, we will discuss the benefits of such an implementation through meaningful data "
              "analysis to gain insight into the performance of Riot Direct Infrastructure.")
doc.add_paragraph(intro_text)

# Player Payload Section
doc.add_heading('Player Payload', level=1)
player_payload_text = ("Due to a lack of information on internet traffic related to VALORANT, the first step was to investigate the IP address for the "
                       "LA California servers. It was later determined that the IP address is universal for all devices despite multiple connections. This "
                       "was confirmed by Riot Games. In this case, the IP was confirmed to be 192.207.0.1.\n\n"
                       "Once the IP was confirmed, an attempt to find what classifies as a player in the internet traffic was made. It was assumed that the player "
                       "would have the most packet data sent to the server, so a count was collected and it was found that the most common payload size is 177 to 127.\n\n"
                       "With a count over 5000, it was assumed that the packet length is the player data. This packet length will be referred to as the player payload "
                       "from here forth.")
doc.add_paragraph(player_payload_text)

# Analysis Section
doc.add_heading('Analysis', level=1)
analysis_text = ("With this, we will use all data within this player payload as the source data for the analysis, specifically the performance of the Riot Direct "
                 "connection by measuring the most important thing in any game: the passage of time. Two games were run during the sample data collection and it was found "
                 "that player identification was consistent in that all packets sent, regardless of their hex code, had all the same header shapes. For instance:\n\n"
                 "Sample 1: 98:3f:ec:57:9c:32:fa:8d\nSample 2: 44:99:a2:b4:46:32:fa:8c\n\n"
                 "These stay consistent with each sample where the raw data changed but the size and the data per sample stayed the same. It may be speculated that this is "
                 "the classifier for that specific player like a UUID (Universal Unique Identifier); however, this is only a theory and no evidence suggests it. But this "
                 "concept of identifying player data through the header has been the most reliable means of filtering data. The data after filtration and after the header will "
                 "be referred to as player data.\n\n"
                 "With data isolated, the meaning behind the player data was investigated; however, the level of entropy suggested a high-level encryption which was expected so "
                 "the investigation was moved toward the first three hex values after the header. According to observations, these involved data synchronization for each packet. "
                 "With UDP this is expected as UDP base packets aren’t transmitted linearly, but this will be elaborated on later.\n\n"
                 "As for the investigation, the graph below indicates the average time difference between each packet. Due to the high level of packets, packets were collected into "
                 "chunks and for each chunk, the average time difference was calculated before plotting.\n\n"
                 "The low time differences show a very fast and reliable network which may be attributed to Riot’s Riot Direct network connection. However, with such speed and the "
                 "nature of UDP, the packets were probably not synchronized. This assumption was surprisingly disproven by the first hex code in the player data which will be referred "
                 "to as the Synchronization ID.\n\n"
                 "Synchronization is defined by the consistent match between the send and receive of packets. UDP tends to trade lower synchronization for speed. To combat this, UDP includes "
                 "a synchronization header that shows when it was sent. This is what the Synchronization ID is referring to, but what differs from the two is that this was plotted based on the "
                 "order in which the packets were exported and no prior sorting was done. This suggests that while the UDP packets would normally be sent at different times due to internet traffic "
                 "congestion, they were sent and received in a perfect order. The spikes and drops are due to the Synchronization ID resetting back to zero.\n\n"
                 "In short, this graph indicates that Riot Direct has optimized the use of UDP in its speed but also combats its downfall which is the large packet loss, so multiple packets are often sent. "
                 "These packet losses can be seen by the separation between the peak/troughs of each spike.\n\n"
                 "In this case, packet loss can be measured by taking the separation in peaks where larger separations indicate the need for more packets to be sent (higher index to make up for packet loss). "
                 "This is noticeable in the graph above with the slight slope change in one of the center spikes while most other slopes seem to be linear. Measuring these differentials will give key insight "
                 "into the packet loss.\n\n"
                 "Similarly, the x-axis was designated as the index to see if a fit with that anomaly in the graph above could be found, where the large spike in the center is the same location as the anomaly in "
                 "the graph before. While these spikes seem large, it’s the drop and rise that matter as each peak is determined by the difference between two peaks in the prior graph.")
doc.add_paragraph(analysis_text)

# Conclusion Section
doc.add_heading('Conclusion', level=1)
conclusion_text = ("Riot Direct is an innovative method of improving competitive integrity. Many steps were taken by their team, one being Riot Direct’s designated internet fiber. This gives a private highway "
                   "for Riot that reduces latency and decreases packet loss. The optimized synchronization is clearly seen and the low packet time difference shows the power that having a designated internet cable "
                   "allows for a better gaming experience, giving them a competitive advantage over other gaming companies, especially in FPS games where every millisecond matters.")
doc.add_paragraph(conclusion_text)

# Save the Document
doc.save("Capstone_Project_Analysis_of_Riot_Direct_in_VALORANT.docx")
